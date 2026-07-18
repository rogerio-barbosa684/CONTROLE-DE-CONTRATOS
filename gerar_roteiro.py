from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# --- styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# --- cover ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Roteiro de Apresentacao\n')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 0xD4, 0xFF)
run = p.add_run('Sistema de Controle de Contratos\n')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x7B, 0x2F, 0xF7)
run = p.add_run('Reuniao com Diretoria')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x66, 0x77, 0x88)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Duracao estimada: ~10 minutos  |  7 slides  |  ~1 minuto por slide')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x55, 0x66, 0x77)

doc.add_paragraph().add_run('').font.size = Pt(6)
doc.add_paragraph().add_run('').font.size = Pt(6)

# --- slide data ---
slides = [
    {
        'num': '01',
        'title': 'Abertura — O Problema',
        'time': '1 min',
        'tone': 'Tom: Direto ao ponto. Mostre que voce entende o problema da empresa. Nao entre em detalhes tecnicos.',
        'script': (
            'Objetivo desta fala: Apresentar o prototipo funcional do sistema de controle de contratos '
            'desenvolvido internamente.\n\n'
            'Fala sugerida:\n'
            '"Bom dia a todos. Hoje quero apresentar um prototipo que desenvolvemos para resolver um problema '
            'que afeta diretamente nossa gestao financeira: o controle descentralizado de contratos e pagamentos.\n\n'
            'Atualmente, usamos planilhas e registros manuais que consomem tempo, geram retrabalho e aumentam o '
            'risco de atrasos. Desenvolvemos um sistema simples, funcional e que ja esta rodando."'
        )
    },
    {
        'num': '02',
        'title': 'O que o sistema faz',
        'time': '1 min 30s',
        'tone': 'Tom: Mostre valor de negocios. Enfatize que substitui planilhas e da visibilidade em tempo real.',
        'script': (
            'Fala sugerida:\n'
            '"O sistema centraliza tres pilares da gestao de contratos:\n\n'
            'Cadastro de contratos — dados completos do fornecedor/cliente, CNPJ/CPF, vigencia, valor e '
            'observacoes. Tudo em um so lugar.\n\n'
            'Parcelas e pagamentos — ao cadastrar um contrato, as parcelas sao geradas automaticamente. Damos '
            'baixa individualmente, com status claro: pendente, pago ou atrasado.\n\n'
            'Dashboard financeiro — visao geral do que esta ativo, vencido, pendente e o total a pagar. '
            'Com filtros por fornecedor, documento e data de vencimento."'
        )
    },
    {
        'num': '03',
        'title': 'Funcionalidades principais',
        'time': '1 min 30s',
        'tone': 'Dica: Se tiver um exemplo real de contrato da empresa, projete o modal de detalhes.',
        'script': (
            'Fala sugerida:\n'
            '"Vou destacar o que considero os diferenciais deste prototipo:\n\n'
            'Busca inteligente — localiza contratos por fornecedor, CPF ou CNPJ. Sem perder tempo procurando.\n\n'
            'Upload de comprovantes — cada pagamento pode ter um comprovante anexado (PDF ou imagem). '
            'Isso e crucial para auditoria.\n\n'
            'Renegociacao de parcelas — quando um contrato e renegociado, as parcelas antigas sao preservadas '
            'e as novas sao geradas automaticamente.\n\n'
            'Modal de detalhes — ao clicar em qualquer contrato, ve todas as informacoes, parcelas, '
            'comprovantes e observacoes."'
        )
    },
    {
        'num': '04',
        'title': 'Modelo de implementacao',
        'time': '1 min 30s',
        'tone': 'Tom: Diretoria se preocupa com custo e risco. Mostre que a versao estatica ja entrega valor '
                'com custo ZERO. A versao servidor e o passo seguinte.',
        'script': (
            'Fala sugerida:\n'
            '"Desenvolvemos duas formas de usar o sistema, porque queriamos flexibilidade:\n\n'
            'Versao estatica (arquivo unico) — R$ 0 de infraestrutura. Basta abrir um arquivo HTML no navegador. '
            'Os dados ficam salvos localmente. Ideal para uso individual ou teste imediato. Nao precisa de '
            'servidor, internet nem instalacao.\n\n'
            'Versao servidor (Flask) — para uso corporativo. Instalada em um servidor com banco de dados SQLite, '
            'acesso via navegador, com autenticacao por usuario e senha. Escalavel para multiplos usuarios na '
            'mesma rede."'
        )
    },
    {
        'num': '05',
        'title': 'Dashboard — demonstracao',
        'time': '1 min 30s',
        'tone': 'Dica: Abra o sistema ao vivo e mostre o dashboard. Filtre por algum fornecedor real para '
                'demonstrar a busca. Isso impressiona mais que slide.',
        'script': (
            'Fala sugerida:\n'
            '"Vou projetar o dashboard para voces verem na pratica.\n\n'
            'Aqui temos:\n'
            'Cards resumo — total de contratos, ativos, vencidos, pagamentos pendentes, em atraso e o valor '
            'total a pagar. Numeros que ajudam a tomar decisoes rapidas.\n\n'
            'Tabela de proximos vencimentos — mostra os 10 pagamentos mais urgentes. Cada linha tem o contrato, '
            'fornecedor, CNPJ/CPF, parcela, data e valor.\n\n'
            'E temos filtros na tela: digite o nome do fornecedor ou CNPJ e a lista ja filtra. Tambem da para '
            'filtrar por data de vencimento."'
        )
    },
    {
        'num': '06',
        'title': 'Beneficios para a empresa',
        'time': '1 min 30s',
        'tone': 'Tom: Diretoria quer saber: economiza quanto? Reduz que risco? Tem custo de implantacao?',
        'script': (
            'Fala sugerida:\n'
            '"O que este sistema entrega de valor concreto:\n\n'
            'Reducao de tempo — o que levava minutos procurando em planilhas agora leva segundos com a busca '
            'integrada.\n\n'
            'Reducao de risco — pagamentos em atraso sao identificados imediatamente no dashboard. Acabou o '
            '"esqueci de pagar".\n\n'
            'Auditabilidade — todo pagamento tem comprovante anexado. Prestacao de contas simplificada.\n\n'
            'Sem dependencia externa — sistema desenvolvido internamente, sem custo de assinatura ou licenca. '
            'Total controle sobre os dados.\n\n'
            'Acessibilidade — funciona em qualquer computador com navegador, sem instalar nada."'
        )
    },
    {
        'num': '07',
        'title': 'Proximos passos e encerramento',
        'time': '1 min 30s',
        'tone': 'Tom: Abra para perguntas. Encerre com tom de colaboracao: o sistema existe, esta funcional, '
                'e agora precisa do direcionamento da diretoria.',
        'script': (
            'Fala sugerida:\n'
            '"Para levar este prototipo a producao, sugiro os seguintes passos:\n\n'
            'Curto prazo (30 dias):\n'
            '— Validacao com usuario piloto (uma area da empresa usando o sistema)\n'
            '— Pequenos ajustes com base no feedback real\n'
            '— Definir se vamos de versao estatica ou servidor\n\n'
            'Medio prazo (60-90 dias):\n'
            '— Implantacao da versao servidor na rede interna\n'
            '— Backup automatico dos dados\n'
            '— Relatorios exportaveis (PDF/Excel)\n\n'
            'Gostaria de ouvir a opniao de voces e sugestoes para priorizar os proximos passos."'
        )
    }
]

# --- write slides ---
for s in slides:
    # header
    p = doc.add_paragraph()
    run = p.add_run(f'Slide {s["num"]}')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0xD4, 0xFF)

    run = p.add_run(f'  {s["title"]}')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    run = p.add_run(f'  ({s["time"]})')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x66, 0x77)

    # script
    p = doc.add_paragraph()
    run = p.add_run(s['script'])
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0xC0, 0xD0, 0xE0)

    # tone/tip
    p = doc.add_paragraph()
    run = p.add_run('💡 ' + s['tone'])
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x7B, 0x2F, 0xF7)

    doc.add_paragraph().add_run('').font.size = Pt(2)

# --- footer ---
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Boa apresentacao!')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x00, 0xD4, 0xFF)

doc.save('roteiro-apresentacao.docx')
print("roteiro-apresentacao.docx gerado com sucesso!")
